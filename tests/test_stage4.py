"""Tests for Stage 4: aggregator, prompts, llm_client, report_builder, docx_renderer, evaluate."""

from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path

import pytest

from src.common.llm_client import (
    DEFAULT_CLIENT,
    HuggingFaceLLMClient,
    LLMClient,
    MockLLMClient,
)
from src.stage4_report_gen.aggregator import (
    compute_contract_risk_score,
    group_by_risk_level,
)
from src.stage4_report_gen.prompts import (
    MAX_CONTRACT_CHARS,
    SYSTEM_PROMPT,
    build_summary_and_conclusion_prompt,
)
from src.stage4_report_gen.report_builder import (
    DISCLAIMER_TEXT,
    assemble_report_dict,
)
from src.stage4_report_gen.docx_renderer import render_docx
from src.stage4_report_gen.evaluate import (
    evaluate_report_completeness,
    evaluate_summaries,
)


# Lightweight clause stand-in. Mirrors both schema variants enough to
# exercise the duck-typed helpers.
@dataclass
class ClauseStub:
    clause_id: str
    document_id: str
    clause_type: str
    clause_text: str
    risk_level: str
    risk_reason: str = ""
    risk_explanation: str = ""
    confidence: float = 0.9


@pytest.fixture
def sample_clauses():
    return [
        ClauseStub("c1", "d1", "Indemnification", "Vendor indemnifies Buyer.",
                   "HIGH", risk_reason="One-sided.", confidence=0.9),
        ClauseStub("c2", "d1", "Termination For Convenience", "30 days notice.",
                   "MEDIUM", risk_reason="Short notice.", confidence=0.8),
        ClauseStub("c3", "d1", "Cap On Liability", "12 months fees.",
                   "LOW", risk_reason="Reasonable.", confidence=0.95),
        ClauseStub("c4", "d1", "Irrevocable Or Perpetual License", "Perpetual.",
                   "HIGH", risk_reason="No revocation.", confidence=0.7),
    ]


# ---------------------------------------------------------------------------
# Aggregator (slimmed to two functions)
# ---------------------------------------------------------------------------

class TestAggregator:

    def test_group_by_risk_level(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        assert {len(groups[k]) for k in ("HIGH", "MEDIUM", "LOW")} == {2, 1, 1}
        assert all(c.risk_level == "HIGH" for c in groups["HIGH"])

    def test_group_normalizes_case_and_missing(self):
        clauses = [
            ClauseStub("a", "d", "X", "x", "high"),       # lowercase
            ClauseStub("b", "d", "X", "x", ""),           # empty → LOW
        ]
        groups = group_by_risk_level(clauses)
        assert len(groups["HIGH"]) == 1
        assert len(groups["LOW"]) == 1

    def test_contract_risk_score_range(self, sample_clauses):
        score = compute_contract_risk_score(sample_clauses)
        assert 0.0 <= score <= 10.0

    def test_contract_risk_score_empty(self):
        assert compute_contract_risk_score([]) == 0.0

    def test_contract_risk_score_all_high_is_max(self):
        clauses = [ClauseStub("a", "d", "X", "x", "HIGH", confidence=1.0)]
        assert compute_contract_risk_score(clauses) == 10.0


# ---------------------------------------------------------------------------
# Prompt builder
# ---------------------------------------------------------------------------

class TestPromptBuilder:

    def test_returns_tuple_of_strings(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        sys_p, usr_p = build_summary_and_conclusion_prompt(
            "This is a contract.", groups,
        )
        assert isinstance(sys_p, str) and isinstance(usr_p, str)
        assert sys_p == SYSTEM_PROMPT
        assert "This is a contract." in usr_p

    def test_includes_risk_counts(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        _, usr = build_summary_and_conclusion_prompt("ct", groups)
        assert "HIGH: 2" in usr
        assert "MEDIUM: 1" in usr
        assert "LOW: 1" in usr

    def test_includes_clause_types_per_tier(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        _, usr = build_summary_and_conclusion_prompt("ct", groups)
        assert "Indemnification" in usr
        assert "Termination For Convenience" in usr

    def test_truncates_long_contracts(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        long_text = "X" * (MAX_CONTRACT_CHARS + 1000)
        _, usr = build_summary_and_conclusion_prompt(long_text, groups)
        assert "[... truncated ...]" in usr
        assert "truncation note" not in usr  # we use a different phrase
        assert "truncated for length" in usr  # the actual phrase

    def test_empty_buckets_handled(self):
        groups = {"HIGH": [], "MEDIUM": [], "LOW": []}
        _, usr = build_summary_and_conclusion_prompt("ct", groups)
        assert "HIGH: 0" in usr
        # No clause-type lines for empty tiers
        assert "HIGH clause types" not in usr


# ---------------------------------------------------------------------------
# LLM client
# ---------------------------------------------------------------------------

class TestLLMClient:

    def test_mock_returns_required_keys(self):
        out = MockLLMClient().generate_json(
            "system", "Risk classification summary — HIGH: 1, MEDIUM: 0, LOW: 0\nContract text:\nThis vendor agreement.",
        )
        assert set(out.keys()) >= {
            "contract_summary", "overall_assessment",
            "high_priority_actions", "medium_priority_actions",
        }
        assert isinstance(out["high_priority_actions"], list)

    def test_mock_uses_risk_signal_from_prompt(self):
        out = MockLLMClient().generate_json(
            "system",
            "Risk classification summary — HIGH: 3, MEDIUM: 2, LOW: 5\n"
            "HIGH clause types: Indemnification, Cap On Liability\n"
            "MEDIUM clause types: Non-Compete\n"
            "Contract text:\nA distribution agreement.",
        )
        # 3 HIGH + 2 MEDIUM + 5 LOW = 10 clauses; mock surfaces this.
        assert "10" in out["contract_summary"]
        # High actions reference the HIGH types from the prompt
        joined = " ".join(out["high_priority_actions"]).lower()
        assert "indemnification" in joined or "cap on liability" in joined

    def test_mock_handles_no_risk_signal(self):
        out = MockLLMClient().generate_json("system", "Contract text:\nFoo.")
        # Still schema-compliant
        assert "contract_summary" in out
        assert isinstance(out["high_priority_actions"], list)

    def test_huggingface_stub_raises(self):
        with pytest.raises(NotImplementedError):
            HuggingFaceLLMClient().generate_json("s", "u")

    def test_default_client_is_mock(self):
        assert isinstance(DEFAULT_CLIENT, LLMClient)
        assert isinstance(DEFAULT_CLIENT, MockLLMClient)


# ---------------------------------------------------------------------------
# Report builder
# ---------------------------------------------------------------------------

class TestReportBuilder:

    def test_assemble_report_dict_shape(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        report = assemble_report_dict(
            document_id="doc_001",
            metadata={"Document Name": "Test", "Parties": "A and B"},
            grouped=groups,
            overall_risk_score=5.5,
            llm_output={
                "contract_summary": "A test contract.",
                "overall_assessment": "Acceptable risk.",
                "high_priority_actions": ["do X"],
                "medium_priority_actions": ["do Y"],
            },
        )
        for key in (
            "document_id", "metadata", "contract_summary", "overall_risk_score",
            "total_clauses", "summary_header", "risk_tables", "conclusion",
            "disclaimer", "generated_at", "models_used",
        ):
            assert key in report
        assert report["document_id"] == "doc_001"
        assert report["total_clauses"] == 4
        assert report["disclaimer"] == DISCLAIMER_TEXT

    def test_no_top_n_cap_on_buckets(self):
        many_high = [
            ClauseStub(f"c{i}", "d", "Indemnification", "x", "HIGH", confidence=0.9)
            for i in range(20)
        ]
        groups = group_by_risk_level(many_high)
        report = assemble_report_dict(
            "doc", {}, groups, 8.0, {},
        )
        assert len(report["risk_tables"]["HIGH"]) == 20

    def test_metadata_canonical_order_with_dashes(self):
        report = assemble_report_dict(
            "doc", {"Parties": "X"}, {"HIGH": [], "MEDIUM": [], "LOW": []},
            0.0, {},
        )
        meta = report["metadata"]
        assert list(meta.keys()) == [
            "Document Name", "Parties", "Agreement Date",
            "Effective Date", "Expiration Date",
        ]
        assert meta["Parties"] == "X"
        assert meta["Document Name"] == "—"

    def test_risk_row_columns(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        report = assemble_report_dict("doc", {}, groups, 5.0, {})
        for tier_rows in report["risk_tables"].values():
            for row in tier_rows:
                assert set(row.keys()) == {
                    "clause_type", "clause_text", "reasoning", "confidence",
                }
                # No recommendation, page_no, content_label
                assert "recommendation" not in row
                assert "page_no" not in row
                assert "content_label" not in row

    def test_llm_output_none_falls_back(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        report = assemble_report_dict("doc", {}, groups, 5.0, None)
        assert report["contract_summary"]  # non-empty fallback
        assert report["conclusion"]["overall_assessment"]


# ---------------------------------------------------------------------------
# DOCX renderer (smoke tests via re-parse)
# ---------------------------------------------------------------------------

class TestDocxRenderer:

    @pytest.fixture
    def base_report(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        return assemble_report_dict(
            document_id="render_test_001",
            metadata={"Document Name": "Test Agreement", "Parties": "A and B"},
            grouped=groups,
            overall_risk_score=5.5,
            llm_output={
                "contract_summary": "This is a test contract summary.",
                "overall_assessment": "Acceptable risk.",
                "high_priority_actions": ["Negotiate indemnification cap."],
                "medium_priority_actions": ["Confirm 30-day notice period."],
            },
        )

    def test_render_creates_docx_file(self, base_report, tmp_path):
        out = render_docx(base_report, output_dir=tmp_path)
        assert out.exists()
        assert out.suffix == ".docx"
        assert out.stem == "render_test_001"

    def test_rendered_docx_is_parseable(self, base_report, tmp_path):
        from docx import Document
        out = render_docx(base_report, output_dir=tmp_path)
        d = Document(str(out))
        # Heading + sections present
        text = "\n".join(p.text for p in d.paragraphs)
        assert "Contract Risk Analysis Report" in text
        assert "Document Metadata" in text
        assert "Contract Summary" in text
        assert "Conclusion & Recommendations" in text
        assert "Disclaimer" in text

    def test_empty_buckets_render_gracefully(self, tmp_path):
        report = assemble_report_dict(
            "empty", {}, {"HIGH": [], "MEDIUM": [], "LOW": []}, 0.0, {},
        )
        out = render_docx(report, output_dir=tmp_path)
        assert out.exists()
        from docx import Document
        d = Document(str(out))
        text = "\n".join(p.text for p in d.paragraphs)
        # Empty-tier note is rendered
        assert "No high-risk clauses identified." in text


# ---------------------------------------------------------------------------
# Evaluate
# ---------------------------------------------------------------------------

class TestEvaluate:

    def test_completeness_passes_for_full_report(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        report = assemble_report_dict(
            "doc", {"Parties": "X"}, groups, 5.0,
            {"contract_summary": "ok", "overall_assessment": "fine",
             "high_priority_actions": [], "medium_priority_actions": []},
        )
        checks = evaluate_report_completeness(report)
        assert all(checks.values()), checks

    def test_completeness_flags_missing_summary(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        report = assemble_report_dict(
            "doc", {"Parties": "X"}, groups, 5.0,
            {"contract_summary": "ok", "overall_assessment": "fine",
             "high_priority_actions": [], "medium_priority_actions": []},
        )
        report["contract_summary"] = ""
        checks = evaluate_report_completeness(report)
        assert checks["has_summary"] is False

    def test_completeness_flags_score_out_of_range(self, sample_clauses):
        groups = group_by_risk_level(sample_clauses)
        report = assemble_report_dict("doc", {}, groups, 5.0, {})
        report["overall_risk_score"] = 99.0
        assert evaluate_report_completeness(report)["score_in_range"] is False

    def test_evaluate_summaries_runs(self):
        result = evaluate_summaries(
            ["the cat sat on the mat"],
            ["a cat sits on the mat"],
        )
        assert set(result.keys()) == {"rouge1", "rouge2", "rougeL"}
        for v in result.values():
            assert 0.0 <= v <= 1.0

    def test_evaluate_summaries_length_mismatch(self):
        with pytest.raises(ValueError):
            evaluate_summaries(["a"], ["b", "c"])
