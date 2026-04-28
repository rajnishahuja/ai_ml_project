"""
Stage 4 DOCX renderer.

Builds the .docx file from the assembled report dict using `python-docx`.
Pure programmatic generation (no Jinja templates) so the entire layout is
auditable in code and there's no template-syncing problem.

Layout (mirrors the spec in the plan):

    Heading: "Contract Risk Analysis Report"
    Sub-line: Document ID, generated timestamp, overall risk score
    Section: Document Metadata (2-column table)
    Section: Contract Summary (paragraph from Mistral)
    Section: HIGH Risk Clauses (4-column table or empty-line note)
    Section: MEDIUM Risk Clauses (same)
    Section: LOW Risk Clauses (same)
    Section: Conclusion & Recommendations
    Section: Disclaimer (italic)

Output path: `<output_dir>/<document_id>.docx`. Returns the path.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


# Risk-tier colors for table headers.
_TIER_COLORS = {
    "HIGH":   RGBColor(0xC0, 0x39, 0x2B),  # red
    "MEDIUM": RGBColor(0xE6, 0x7E, 0x22),  # orange
    "LOW":    RGBColor(0x27, 0xAE, 0x60),  # green
}


# ---------------------------------------------------------------------------
# Small helpers (kept private — caller uses render_docx)
# ---------------------------------------------------------------------------

def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_metadata_table(doc: Document, metadata: dict[str, str]) -> None:
    """Two-column table: field name, value. One row per metadata field."""
    table = doc.add_table(rows=len(metadata), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (field, value) in enumerate(metadata.items()):
        name_cell = table.rows[i].cells[0]
        value_cell = table.rows[i].cells[1]
        name_cell.text = field
        value_cell.text = str(value) if value else "—"
        # Bold the field name
        for run in name_cell.paragraphs[0].runs:
            run.bold = True

    # Column widths — narrower label column.
    for row in table.rows:
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)


def _add_risk_table(
    doc: Document,
    tier: str,
    rows: list[dict[str, Any]],
) -> None:
    """One risk-tier section: heading + table (or empty-line note).

    Columns: Clause Type | Clause Text | Reasoning | Confidence
    """
    color = _TIER_COLORS.get(tier, RGBColor(0, 0, 0))
    h = doc.add_heading(level=2)
    run = h.add_run(f"{tier} Risk Clauses (n={len(rows)})")
    run.font.color.rgb = color

    if not rows:
        para = doc.add_paragraph(f"No {tier.lower()}-risk clauses identified.")
        para.runs[0].italic = True
        return

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    headers = ["Clause Type", "Clause Text", "Reasoning", "Confidence"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        cells[0].text = row.get("clause_type", "")
        cells[1].text = row.get("clause_text", "")
        cells[2].text = row.get("reasoning", "")
        conf = row.get("confidence")
        cells[3].text = f"{conf:.2f}" if conf is not None else "—"
        for cell in cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Tuned column widths so clause text gets the most room.
    widths = [Cm(3.5), Cm(7.0), Cm(5.0), Cm(2.0)]
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def _add_conclusion(doc: Document, conclusion: dict[str, Any]) -> None:
    _add_heading(doc, "Conclusion & Recommendations", level=1)
    doc.add_paragraph(
        conclusion.get("overall_assessment") or "(No overall assessment provided.)"
    )

    high = conclusion.get("high_priority_actions") or []
    if high:
        _add_heading(doc, "High-Priority Actions", level=2)
        for action in high:
            doc.add_paragraph(action, style="List Bullet")

    medium = conclusion.get("medium_priority_actions") or []
    if medium:
        _add_heading(doc, "Medium-Priority Actions", level=2)
        for action in medium:
            doc.add_paragraph(action, style="List Bullet")


def _add_disclaimer(doc: Document, text: str) -> None:
    _add_heading(doc, "Disclaimer", level=2)
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def render_docx(report: dict[str, Any], output_dir: Path | str) -> Path:
    """Render the report dict to a .docx file.

    Args:
        report: Output of `assemble_report_dict(...)`.
        output_dir: Directory to write into. Created if missing.

    Returns:
        Path to the written `.docx` file: `<output_dir>/<document_id>.docx`.
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    document_id = str(report.get("document_id") or "unknown")
    out_path = out_dir / f"{document_id}.docx"

    doc = Document()

    # ── Title block ─────────────────────────────────────────────────────
    title = doc.add_heading("Contract Risk Analysis Report", level=0)
    sub = doc.add_paragraph()
    sub.add_run(f"Document ID: ").bold = True
    sub.add_run(f"{document_id}\n")
    sub.add_run("Generated: ").bold = True
    sub.add_run(f"{report.get('generated_at', '—')}\n")
    sub.add_run("Overall Risk Score: ").bold = True
    score = report.get("overall_risk_score")
    sub.add_run(f"{score:.2f} / 10" if isinstance(score, (int, float)) else "—")

    # ── Document metadata ───────────────────────────────────────────────
    _add_heading(doc, "Document Metadata", level=1)
    _add_metadata_table(doc, report.get("metadata", {}))

    # ── Contract summary (Mistral-generated) ────────────────────────────
    _add_heading(doc, "Contract Summary", level=1)
    doc.add_paragraph(report.get("contract_summary", ""))

    # ── Three risk tables ───────────────────────────────────────────────
    _add_heading(doc, "Risk-Classified Clauses", level=1)
    risk_tables = report.get("risk_tables", {})
    for tier in ("HIGH", "MEDIUM", "LOW"):
        _add_risk_table(doc, tier, risk_tables.get(tier, []))

    # ── Conclusion ──────────────────────────────────────────────────────
    _add_conclusion(doc, report.get("conclusion", {}))

    # ── Disclaimer ──────────────────────────────────────────────────────
    _add_disclaimer(doc, report.get("disclaimer", ""))

    doc.save(out_path)
    logger.info("Wrote DOCX report to %s", out_path)
    return out_path
